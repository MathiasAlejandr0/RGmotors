# -*- coding: utf-8 -*-
"""Genera documentos Word profesionales del proyecto RG Motors Puerto Montt."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from datetime import date
import os

OUT = os.path.join(os.path.dirname(os.path.dirname(__file__)), "documentos")
os.makedirs(OUT, exist_ok=True)

NAVY = RGBColor(0x1A, 0x36, 0x5D)
ACCENT = RGBColor(0x2B, 0x6C, 0xB0)
MUTED = RGBColor(0x4A, 0x55, 0x68)
LINE = RGBColor(0xC5, 0xCD, 0xD8)
BODY_FONT = "Calibri"
TITLE_FONT = "Calibri Light"


def set_run_font(run, size=11, bold=False, color=None, font=BODY_FONT, italic=False):
    run.font.name = font
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def setup_doc(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)

    styles = doc.styles

    # Normal
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15
    pf.space_after = Pt(8)
    pf.space_before = Pt(0)

    for level, size, before, after in [
        (1, 16, 18, 10),
        (2, 13, 14, 8),
        (3, 12, 12, 6),
    ]:
        style = styles[f"Heading {level}"]
        style.font.name = TITLE_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = NAVY
        style._element.rPr.rFonts.set(qn("w:ascii"), TITLE_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), TITLE_FONT)
        sp = style.paragraph_format
        sp.space_before = Pt(before)
        sp.space_after = Pt(after)
        sp.line_spacing = 1.15
        sp.keep_with_next = True


def add_page_number(paragraph):
    run = paragraph.add_run()
    set_run_font(run, size=9, color=MUTED)
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def add_footer(doc, code):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{code}  ·  Confidencial  ·  Página ")
    set_run_font(run, size=9, color=MUTED)
    add_page_number(p)


def horizontal_line(doc, color="1A365D"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def thin_line(doc):
    return horizontal_line(doc, "C5CDD8")


def page_break(doc):
    doc.add_page_break()


def add_para(
    doc,
    text,
    bold=False,
    size=11,
    align="left",
    space_after=8,
    space_before=0,
    color=None,
    font=BODY_FONT,
    italic=False,
):
    p = doc.add_paragraph()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color, font=font, italic=italic)
    return p


def add_heading_styled(doc, text, level=1):
    return doc.add_heading(text, level=level)


def shade_cell(cell, hex_color="1A365D"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=10, center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(
            cell, h, bold=True, color=RGBColor(255, 255, 255), size=10, center=True
        )
        shade_cell(cell, "1A365D")
    for r_i, row in enumerate(rows):
        bg = "F7F9FC" if r_i % 2 == 0 else "FFFFFF"
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            set_cell_text(cell, str(val), size=9)
            shade_cell(cell, bg)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(10)
    return table


def cb(label):
    return f"☐  {label}"


def cover_page(
    doc,
    doc_type,
    title,
    subtitle,
    code,
    version="1.0",
):
    # Espacio superior
    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    add_para(
        doc,
        "RG MOTORS",
        bold=True,
        size=12,
        align="center",
        color=ACCENT,
        font=TITLE_FONT,
        space_after=4,
    )
    add_para(
        doc,
        "SUCURSAL PUERTO MONTT",
        bold=False,
        size=11,
        align="center",
        color=MUTED,
        space_after=18,
    )
    horizontal_line(doc, "1A365D")

    add_para(
        doc,
        doc_type.upper(),
        bold=False,
        size=11,
        align="center",
        color=ACCENT,
        font=TITLE_FONT,
        space_after=14,
        space_before=18,
    )
    add_para(
        doc,
        title,
        bold=True,
        size=22,
        align="center",
        color=NAVY,
        font=TITLE_FONT,
        space_after=12,
    )
    add_para(
        doc,
        subtitle,
        size=12,
        align="center",
        color=MUTED,
        space_after=22,
    )
    thin_line(doc)

    add_para(doc, "", space_after=8)
    meta = [
        ("Código", code),
        ("Versión", version),
        ("Fecha", date.today().strftime("%d de %B de %Y").replace(
            "January", "enero"
        ).replace("February", "febrero")
         .replace("March", "marzo").replace("April", "abril")
         .replace("May", "mayo").replace("June", "junio")
         .replace("July", "julio").replace("August", "agosto")
         .replace("September", "septiembre").replace("October", "octubre")
         .replace("November", "noviembre").replace("December", "diciembre")),
        ("Clasificación", "Uso interno — Confidencial"),
        ("Alcance", "Sucursal Puerto Montt"),
    ]
    # Fecha simple sin locale issues
    meta[2] = ("Fecha", date.today().strftime("%d/%m/%Y"))

    table = doc.add_table(rows=len(meta), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(meta):
        set_cell_text(table.rows[i].cells[0], k, bold=True, size=10, color=NAVY)
        set_cell_text(table.rows[i].cells[1], v, size=10, color=MUTED)
        table.rows[i].cells[0].width = Cm(4.5)
        table.rows[i].cells[1].width = Cm(9)
        shade_cell(table.rows[i].cells[0], "F0F4F8")
        shade_cell(table.rows[i].cells[1], "FFFFFF")

    for _ in range(4):
        doc.add_paragraph()

    add_para(
        doc,
        "Plataforma web de venta de vehículos usados\ncon catálogo digital, tour 360° y administración autónoma",
        size=10,
        align="center",
        color=MUTED,
        italic=True,
        space_after=6,
    )
    horizontal_line(doc, "C5CDD8")
    add_para(
        doc,
        "Documento elaborado en el marco del proyecto informático de implementación digital",
        size=9,
        align="center",
        color=MUTED,
        space_after=0,
    )
    page_break(doc)


def toc_page(doc, items, title="Índice de contenidos"):
    """items: list of (number_label, title, level) level 1 or 2"""
    add_para(
        doc,
        title,
        bold=True,
        size=18,
        align="left",
        color=NAVY,
        font=TITLE_FONT,
        space_after=6,
    )
    thin_line(doc)
    add_para(
        doc,
        "Seleccione un tema para ubicar rápidamente la sección correspondiente en el documento.",
        size=10,
        color=MUTED,
        italic=True,
        space_after=16,
    )

    for label, text, level in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3 if level == 1 else 1)
        p.paragraph_format.space_after = Pt(3 if level == 1 else 1)
        p.paragraph_format.line_spacing = 1.2
        if level == 2:
            p.paragraph_format.left_indent = Cm(0.75)

        run_num = p.add_run(f"{label}  ")
        set_run_font(
            run_num,
            size=11 if level == 1 else 10,
            bold=level == 1,
            color=ACCENT if level == 1 else MUTED,
            font=BODY_FONT,
        )
        run_txt = p.add_run(text)
        set_run_font(
            run_txt,
            size=11 if level == 1 else 10,
            bold=level == 1,
            color=NAVY if level == 1 else MUTED,
            font=BODY_FONT,
        )

        # leader dots via tab isn't reliable; use spaced dots
        # skip visual page numbers for manual TOC (Word won't auto-update without TOC field)

    page_break(doc)


def section_intro(doc, text):
    add_para(doc, text, align="justify", size=11, space_after=10, color=MUTED, italic=True)


# ---------------------------------------------------------------------------
# DOC 1
# ---------------------------------------------------------------------------
def build_solicitud():
    doc = Document()
    setup_doc(doc)
    add_footer(doc, "RG-PM-SOL-001")

    cover_page(
        doc,
        doc_type="Solicitud formal",
        title="Accesos, decisiones\ny aprobaciones",
        subtitle="Insumos, infraestructura e integraciones comerciales\nrequeridos para continuar el proyecto",
        code="RG-PM-SOL-001",
    )

    toc_page(
        doc,
        [
            ("1.", "Objetivo del documento", 1),
            ("2.", "Alcance confirmado", 1),
            ("3.", "Solicitud de fotografías y videos del stock", 1),
            ("3.1", "Medio de entrega", 2),
            ("3.2", "Estructura de carpetas sugerida", 2),
            ("3.3", "Formatos y calidad mínima", 2),
            ("3.4", "Datos asociados a cada vehículo", 2),
            ("4.", "Compra y gestión del dominio (.cl)", 1),
            ("5.", "Elección de base de datos", 1),
            ("6.", "Elección de hosting", 1),
            ("7.", "Integraciones comerciales: WebPay y financiamiento", 1),
            ("7.1", "WebPay (Transbank)", 2),
            ("7.2", "Financiamiento (Forum u otras)", 2),
            ("8.", "Principios de gobernanza", 1),
            ("9.", "Plazo de respuesta solicitado", 1),
            ("10.", "Aprobaciones", 1),
        ],
    )

    add_heading_styled(doc, "1. Objetivo del documento", 1)
    section_intro(
        doc,
        "Este apartado define el propósito de la solicitud y el marco de decisión que se espera de la Empresa.",
    )
    add_para(
        doc,
        "El presente documento tiene por objeto formalizar las solicitudes y decisiones que la Empresa debe adoptar "
        "para continuar el desarrollo del sitio web de la sucursal Puerto Montt. Se busca evitar ambigüedades, "
        "retrasos por falta de insumos y conflictos de responsabilidad sobre costos, accesos y contratación de servicios.",
        align="justify",
    )

    add_heading_styled(doc, "2. Alcance confirmado", 1)
    add_para(doc, "Para efectos de viabilidad del proyecto, el alcance de la primera versión (v1) queda definido como:")
    for t in [
        "Sitio web orientado exclusivamente a la sucursal de Puerto Montt.",
        "Catálogo de vehículos, fichas de detalle, tour 360°, contacto comercial y panel de administración.",
        "Fuera de alcance v1: otras sucursales, CRM interno completo, app móvil nativa.",
    ]:
        add_para(doc, f"• {t}", space_after=4)

    add_heading_styled(doc, "3. Solicitud de fotografías y videos del stock", 1)
    add_para(
        doc,
        "Se requiere acceso a material audiovisual de los vehículos existentes en Puerto Montt. "
        "Sin este insumo no es posible publicar un catálogo profesional ni tours 360° reales.",
        align="justify",
    )

    add_heading_styled(doc, "3.1. Medio de entrega (elige la Empresa)", 2)
    add_para(
        doc,
        "Se recomienda un canal que preserve la calidad original (sin compresión agresiva de WhatsApp). "
        "Marque una opción:",
        align="justify",
    )
    for t in [
        cb("Google Drive (recomendado) — carpeta compartida"),
        cb("Microsoft OneDrive / SharePoint"),
        cb("Dropbox"),
        cb("WeTransfer / Smash (lotes puntuales de alta calidad)"),
        cb("Otro: ________________________________"),
    ]:
        add_para(doc, t, space_after=4)
    add_para(
        doc,
        "Importante: WhatsApp comprime fotos y videos. Puede usarse solo para coordinación, no como archivo maestro.",
        bold=True,
        size=10,
        color=NAVY,
    )

    add_heading_styled(doc, "3.2. Estructura de carpetas sugerida", 2)
    add_para(
        doc,
        "Una carpeta por vehículo, con nombre: MARCA_MODELO_AÑO_PATENTEINTERNA "
        "(ejemplo: TOYOTA_RAV4_2021_ABCD12)",
    )
    for t in [
        "/fotos/ → 8 a 12 imágenes JPG/PNG (frente, atras, laterales, interior, motor, detalle)",
        "/video_360/ → 1 video de vuelta completa (30–60 s, horizontal, buena luz)",
        "/hero/ → 1 foto principal para listado del catálogo",
    ]:
        add_para(doc, f"• {t}", space_after=4)

    add_heading_styled(doc, "3.3. Formatos y calidad mínima", 2)
    add_table(
        doc,
        ["Tipo", "Formato preferido", "Recomendación"],
        [
            ["Fotos", "JPG o PNG", "Mínimo 1920 px en el lado largo; sin filtros extremos"],
            ["Video 360", "MP4 (H.264) o MOV", "Full HD (1080p) o superior; no enviar por WhatsApp"],
            ["Peso", "Sin límite artificial", "Mejor archivo pesado y nítido que liviano y pixelado"],
        ],
    )

    add_heading_styled(doc, "3.4. Datos asociados a cada vehículo (planilla)", 2)
    add_para(doc, "Además del material, se solicita planilla (Excel/Google Sheets) con:")
    for t in [
        "Marca, modelo, año, precio, kilometraje, combustible, transmisión, color",
        "Estado: disponible / reservado / vendido",
        "Patente (uso interno; no se publicará salvo instrucción expresa)",
        "Observaciones comerciales (opcional)",
    ]:
        add_para(doc, f"• {t}", space_after=4)

    add_heading_styled(doc, "4. Compra y gestión del dominio (.cl)", 1)
    add_para(
        doc,
        "El dominio rgmotors.cl ya se encuentra registrado por terceros. La Empresa debe proponer y aprobar "
        "un nombre disponible en NIC Chile (www.nic.cl). El dominio debe quedar a nombre del RUT de la Empresa.",
        align="justify",
    )
    add_table(
        doc,
        ["Concepto", "Detalle"],
        [
            ["Registrador", "NIC Chile (oficial .cl)"],
            ["Costo referencial", "$9.990 CLP por año (exento de IVA, tarifa NIC Chile vigente)"],
            ["Titular", "Razón social / RUT de la Empresa (no cuenta personal del desarrollador)"],
            ["Responsable interno", "________________________________"],
        ],
    )
    add_para(doc, "Propuestas de nombre (completar y marcar disponibilidad en nic.cl):")
    for t in [
        cb("Opción 1: ________________________.cl"),
        cb("Opción 2: ________________________.cl"),
        cb("Opción 3: ________________________.cl"),
    ]:
        add_para(doc, t, space_after=4)
    add_para(
        doc,
        "Nombre aprobado por la Empresa: ________________________.cl     Fecha: ____/____/________",
        space_after=12,
    )

    add_heading_styled(doc, "5. Elección de base de datos / backend de datos", 1)
    add_para(
        doc,
        "Se requiere un servicio en la nube para almacenar el catálogo, imágenes y usuarios del panel de administración. "
        "A continuación se presentan alternativas comparadas para decisión formal.",
        align="justify",
    )
    add_table(
        doc,
        ["Criterio", "Supabase (recomendada)", "Firebase (Google)", "Solo archivos en el código"],
        [
            ["Tipo", "PostgreSQL + Auth + Storage", "NoSQL + Auth + Storage", "Sin base real"],
            ["Costo inicio", "Free; Pro ~US$25/mes", "Pay-as-you-go", "US$0"],
            ["Panel admin", "Excelente para catálogo", "Posible, menos natural p/ filtros", "Requiere programador"],
            ["Fotos/videos", "Storage incluido", "Storage incluido", "GitHub (no ideal)"],
            ["Apto producción", "Sí (Plan Pro)", "Sí", "Solo demo"],
            ["Autonomía Empresa", "Alta (con admin web)", "Media-Alta", "Nula"],
        ],
    )
    add_para(doc, "Recomendación técnica: Supabase Plan Pro para producción.", bold=True, size=10, color=NAVY)
    add_para(doc, "Decisión de la Empresa:")
    for t in [cb("Supabase (recomendado)"), cb("Firebase"), cb("Otra: ________________________")]:
        add_para(doc, t, space_after=4)
    add_para(doc, "Cuenta/correo corporativo para alta del servicio: ________________________")
    add_para(doc, "Titular de facturación: ☐ Empresa   ☐ Otro: ____________")

    add_heading_styled(doc, "6. Elección de hosting (alojamiento web)", 1)
    add_para(
        doc,
        "GitHub Pages se utiliza únicamente como demo. Para operación profesional se requiere hosting pagado "
        "compatible con Next.js y administración autónoma.",
        align="justify",
    )
    add_table(
        doc,
        ["Criterio", "Vercel (recomendado)", "Cloudflare Pages", "VPS (DigitalOcean/etc.)"],
        [
            ["Encaje con Next.js", "Óptimo", "Bueno (con límites)", "Manual"],
            ["Costo referencial", "Pro ~US$20/mes", "Free / planes pagos", "Desde ~US$6–24/mes + mantención"],
            ["Deploy", "Automático desde GitHub", "Automático", "Configuración y mantención propia"],
            ["SSL / HTTPS", "Incluido", "Incluido", "Configurar"],
            ["Esfuerzo interno", "Bajo", "Medio", "Alto (necesita soporte TI)"],
            ["Ideal si…", "Quieren simplicidad y velocidad", "Priorizan CDN global", "Tienen administrador de servidores"],
        ],
    )
    add_para(
        doc,
        "Costo mensual estimado stack recomendado (Vercel Pro + Supabase Pro): ~US$45/mes + dominio anual.",
        size=10,
        color=MUTED,
    )
    add_para(doc, "Decisión de la Empresa:")
    for t in [
        cb("Vercel Pro (recomendado)"),
        cb("Cloudflare Pages"),
        cb("VPS / otro: ________________________"),
    ]:
        add_para(doc, t, space_after=4)
    add_para(doc, "Cuenta/correo corporativo para alta: ________________________")

    add_heading_styled(doc, "7. Integraciones comerciales: WebPay y financiamiento", 1)
    add_para(
        doc,
        "Estas integraciones dependen de contratos comerciales de la Empresa con terceros. "
        "El desarrollo técnico solo puede iniciarse cuando existan credenciales y canales oficiales.",
        align="justify",
    )

    add_heading_styled(doc, "7.1. WebPay (Transbank)", 2)
    add_para(doc, "Requisitos habituales (responsabilidad de la Empresa):")
    for t in [
        "Empresa formalizada con inicio de actividades en SII y giro coherente.",
        "Cuenta corriente a nombre del comercio.",
        "Sitio con dominio propio, HTTPS, datos del comercio, políticas y términos.",
        "Contratación/afiliación en Transbank y certificación técnica del flujo de pago.",
    ]:
        add_para(doc, f"• {t}", space_after=4)
    add_para(doc, "Documentación técnica: https://www.transbankdevelopers.cl", size=10, color=ACCENT)
    add_para(doc, "Solicitud a la Empresa:")
    for t in [
        cb("Autorizamos gestionar afiliación WebPay Plus con Transbank"),
        cb("Entregaremos código de comercio y credenciales de integración/producción"),
        cb("Designamos contacto interno Transbank: ________________________ / Tel: ____________"),
        cb("Por ahora NO incluimos pagos online en v1 (solo contacto/WhatsApp)"),
    ]:
        add_para(doc, t, space_after=4)

    add_heading_styled(doc, "7.2. Financiamiento (Forum u otras financieras)", 2)
    add_para(
        doc,
        "La conexión con Forum u otras entidades de crédito automotriz no es un plugin genérico. "
        "Depende del acuerdo comercial vigente de la automotora con la financiera y del canal que ellos habiliten "
        "(enlace a cotizador, integración dealer, API, proceso asistido por ejecutivo, etc.).",
        align="justify",
    )
    add_para(doc, "Se solicita a la Empresa:")
    for t in [
        cb("Indicar financiera(s) preferente(s): Forum / Otra: ____________"),
        cb("Entregar contacto del ejecutivo comercial de la financiera"),
        cb("Consultar formalmente si existe API, link de cotización, iframe o flujo dealer"),
        cb("Mientras no exista canal técnico: botón “Simular financiamiento” → WhatsApp / formulario / link externo"),
    ]:
        add_para(doc, t, space_after=4)
    add_para(doc, "Ejecutivo financiera: ________________________  Correo/Tel: ________________________")

    add_heading_styled(doc, "8. Principios de gobernanza (para evitar conflictos de carga)", 1)
    for t in [
        "Las cuentas de dominio, hosting, base de datos y pasarelas deben quedar a nombre de la Empresa.",
        "El desarrollador no será titular ni responsable de pagos recurrentes de servicios cloud.",
        "Los contenidos (fotos, precios, stock) son responsabilidad comercial de la sucursal.",
        "Los cambios de alcance se formalizarán por escrito (correo o acta) antes de ejecutarse.",
        "Los plazos de 3 meses asumen respuesta oportuna a las solicitudes de este documento.",
    ]:
        add_para(doc, f"• {t}", space_after=4)

    add_heading_styled(doc, "9. Plazo de respuesta solicitado", 1)
    add_para(
        doc,
        "Se solicita devolver este documento firmado (o responder por correo con las casillas marcadas) "
        "en un plazo máximo de 10 días hábiles, a fin de no detener el cronograma del proyecto.",
        align="justify",
    )

    add_heading_styled(doc, "10. Aprobaciones", 1)
    add_table(
        doc,
        ["Rol", "Nombre", "Firma", "Fecha"],
        [
            ["Solicitante / Responsable técnico", "", "", ""],
            ["Representante Empresa / Sucursal Puerto Montt", "", "", ""],
            ["Responsable comercial / gerencia", "", "", ""],
        ],
        col_widths=[5, 4, 4, 3],
    )
    add_para(doc, "Observaciones de la Empresa:", bold=True, color=NAVY)
    add_para(doc, "_" * 78, color=LINE, space_after=10)
    add_para(doc, "_" * 78, color=LINE, space_after=10)
    add_para(doc, "_" * 78, color=LINE, space_after=10)

    path = os.path.join(OUT, "01-Solicitud-Decisiones-y-Accesos-RGMotors-PuertoMontt.docx")
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# DOC 2
# ---------------------------------------------------------------------------
def build_proyecto():
    doc = Document()
    setup_doc(doc)
    add_footer(doc, "RG-PM-PRY-001")

    cover_page(
        doc,
        doc_type="Documento de proyecto informático",
        title="Plataforma web\nRG Motors Puerto Montt",
        subtitle="Catálogo digital, tour 360°, administración autónoma\ne integraciones comerciales",
        code="RG-PM-PRY-001",
    )

    toc_page(
        doc,
        [
            ("1.", "Resumen ejecutivo", 1),
            ("2.", "Antecedentes y necesidad de la empresa", 1),
            ("2.1", "Contexto", 2),
            ("2.2", "Problema", 2),
            ("2.3", "Necesidad planteada", 2),
            ("3.", "Objetivos", 1),
            ("4.", "Alcance del proyecto", 1),
            ("5.", "Estudio de mercado (síntesis)", 1),
            ("6.", "Stakeholders y responsabilidades", 1),
            ("7.", "Arquitectura y tecnologías", 1),
            ("8.", "Módulos funcionales", 1),
            ("9.", "Costos del proyecto", 1),
            ("10.", "Cronograma de producción (3 meses)", 1),
            ("11.", "Riesgos y mitigaciones", 1),
            ("12.", "Plan de desligue / continuidad operativa", 1),
            ("13.", "Criterios de aceptación (v1)", 1),
            ("14.", "Aspectos legales y de datos", 1),
            ("15.", "Anexos", 1),
            ("16.", "Aprobación del documento", 1),
        ],
    )

    add_heading_styled(doc, "1. Resumen ejecutivo", 1)
    section_intro(doc, "Visión general del proyecto, su propósito y el horizonte de ejecución.")
    add_para(
        doc,
        "El presente documento describe el proyecto de desarrollo e implementación de una plataforma web "
        "para la sucursal de Puerto Montt de RG Motors. La solución permitirá publicar el stock de vehículos "
        "usados con información comercial clara, experiencia visual diferenciadora (tour 360°), canales de "
        "contacto y, en etapas posteriores, integración con medios de pago (WebPay) y financiamiento automotriz "
        "(Forum u otros). El horizonte de ejecución considerado es de tres (3) meses, alineado a una práctica "
        "profesional / proyecto de implementación controlada.",
        align="justify",
    )

    add_heading_styled(doc, "2. Antecedentes y necesidad de la empresa", 1)
    add_heading_styled(doc, "2.1. Contexto", 2)
    add_para(
        doc,
        "La comercialización de vehículos usados en Chile es altamente competitiva y digitalmente mediada: "
        "el cliente investiga en línea antes de visitar la sucursal. Una presencia web débil o desactualizada "
        "reduce la captación de leads y obliga a depender de portales de terceros o de difusión informal.",
        align="justify",
    )
    add_heading_styled(doc, "2.2. Problema", 2)
    for t in [
        "Falta de un canal digital propio, profesional y actualizable por la sucursal.",
        "Dificultad para mostrar el vehículo de forma confiable a distancia (fotos limitadas / sin 360°).",
        "Dependencia de personas técnicas para cada cambio de stock o precio.",
        "Necesidad futura de pagos en línea y derivación a financiamiento.",
    ]:
        add_para(doc, f"• {t}", space_after=4)
    add_heading_styled(doc, "2.3. Necesidad planteada", 2)
    add_para(
        doc,
        "Contar con un sitio web propio de la sucursal Puerto Montt, autoadministrable, con catálogo real, "
        "tour 360°, contacto comercial y base preparada para WebPay y financieras, quedando las cuentas "
        "y la operación a nombre de la Empresa para continuidad operativa.",
        align="justify",
    )

    add_heading_styled(doc, "3. Objetivos", 1)
    add_heading_styled(doc, "3.1. Objetivo general", 2)
    add_para(
        doc,
        "Diseñar, desarrollar e implementar una plataforma web profesional para la venta de vehículos usados "
        "de la sucursal Puerto Montt, con administración autónoma y despliegue en infraestructura pagada.",
        align="justify",
    )
    add_heading_styled(doc, "3.2. Objetivos específicos", 2)
    for t in [
        "Publicar catálogo de stock con fichas detalladas y medios audiovisuales de calidad.",
        "Implementar visor 360° a partir de video/fotos de cada unidad.",
        "Entregar panel de administración para alta/edición/baja de vehículos sin modificar código.",
        "Configurar dominio .cl, hosting y base de datos a nombre de la Empresa.",
        "Dejar preparada la integración con WebPay y el canal de financiamiento (Forum u otro).",
        "Documentar la entrega para desligue técnico del desarrollador.",
    ]:
        add_para(doc, f"• {t}", space_after=4)

    add_heading_styled(doc, "4. Alcance del proyecto", 1)
    add_heading_styled(doc, "4.1. Incluido en v1 (3 meses)", 2)
    for t in [
        "Sitio público: inicio, catálogo, ficha de vehículo, tour 360°, contacto (WhatsApp/formulario).",
        "Panel admin: CRUD de vehículos, carga de fotos, estado del stock.",
        "Optimización de carga del 360° para uso en celular.",
        "Despliegue en hosting profesional + dominio .cl.",
        "Capacitación breve y manual de operación.",
    ]:
        add_para(doc, f"• {t}", space_after=4)
    add_heading_styled(doc, "4.2. Fuera de alcance v1 (posterior)", 2)
    for t in [
        "Multi-sucursal / inventario federado nacional.",
        "App móvil nativa.",
        "ERP/contabilidad completa.",
        "Integración avanzada con portales (Chileautos, etc.) salvo acuerdo explícito.",
        "WebPay y Forum en producción solo si la Empresa entrega contratos y credenciales a tiempo.",
    ]:
        add_para(doc, f"• {t}", space_after=4)

    add_heading_styled(doc, "5. Estudio de mercado (síntesis)", 1)
    add_heading_styled(doc, "5.1. Mercado objetivo", 2)
    add_para(
        doc,
        "Compradores particulares y empresas en la zona sur (foco Puerto Montt y región de Los Lagos) "
        "que buscan vehículos usados con información transparente, respaldo local y opciones de financiamiento.",
        align="justify",
    )
    add_heading_styled(doc, "5.2. Tendencias relevantes", 2)
    for t in [
        "El journey de compra automotriz inicia online (búsqueda, comparación, contacto).",
        "La confianza aumenta con evidencia visual rica (galería completa / 360° / video).",
        "Los portales horizontales concentran demanda, pero la automotora necesita marca propia y captura de lead directa.",
        "El financiamiento es parte crítica de la conversión; la web debe facilitar la derivación.",
    ]:
        add_para(doc, f"• {t}", space_after=4)
    add_heading_styled(doc, "5.3. Competencia digital (tipos)", 2)
    add_table(
        doc,
        ["Tipo", "Fortaleza", "Oportunidad para RG Motors PM"],
        [
            ["Portales nacionales", "Alto tráfico", "Marca propia + atención local + 360° diferencial"],
            ["Sitios de automotoras", "Presencia de marca", "Admin simple + carga rápida móvil + financiamiento visible"],
            ["Redes sociales", "Alcance barato", "Derivar a ficha web completa y stock actualizado"],
        ],
    )
    add_heading_styled(doc, "5.4. Propuesta de valor digital", 2)
    add_para(
        doc,
        "Una vitrina propia de Puerto Montt, actualizable por el equipo comercial, con experiencia 360° "
        "y camino claro a contacto y financiamiento, reduciendo fricción entre interés online y visita a sucursal.",
        align="justify",
    )

    add_heading_styled(doc, "6. Stakeholders y responsabilidades", 1)
    add_table(
        doc,
        ["Actor", "Responsabilidad principal"],
        [
            ["Empresa / Sucursal Puerto Montt", "Aprobar alcance, comprar dominio, pagar cloud, entregar stock/fotos, contratos WebPay/Forum"],
            ["Responsable técnico del proyecto", "Diseño, desarrollo, despliegue inicial, documentación y capacitación"],
            ["Transbank", "Afiliación, certificación y operación de WebPay"],
            ["Forum u otra financiera", "Definir canal de integración / cotización para dealers"],
            ["Proveedores cloud", "Disponibilidad de hosting, DB y storage (Vercel, Supabase, etc.)"],
        ],
    )

    add_heading_styled(doc, "7. Arquitectura y tecnologías", 1)
    add_heading_styled(doc, "7.1. Stack recomendado", 2)
    add_table(
        doc,
        ["Capa", "Tecnología", "Motivo"],
        [
            ["Frontend / App", "Next.js (React) + TypeScript", "Prototipo existente; SEO, rendimiento y APIs"],
            ["UI", "Tailwind CSS", "Velocidad de desarrollo y consistencia visual"],
            ["Tour 360°", "Visor propio + procesamiento de video a frames", "Diferenciación comercial"],
            ["Base de datos", "Supabase (PostgreSQL)", "Datos, auth admin y storage en un solo servicio"],
            ["Hosting", "Vercel", "Deploy automático, SSL, ideal Next.js"],
            ["Dominio", "NIC Chile (.cl)", "Identidad local y práctica estándar en Chile"],
            ["Pagos (etapa 2/3)", "WebPay Plus (Transbank)", "Estándar de mercado nacional"],
            ["Financiamiento", "Forum / link o API según acuerdo", "Canal comercial realista"],
        ],
    )
    add_heading_styled(doc, "7.2. Alternativas evaluadas", 2)
    add_para(doc, "Base de datos: Supabase (recomendada) vs Firebase vs mantener datos en código (solo demo).")
    add_para(doc, "Hosting: Vercel (recomendado) vs Cloudflare Pages vs VPS autogestionado.")
    add_para(doc, "Detalle de costos y decisión formal: ver documento RG-PM-SOL-001.", size=10, color=MUTED)

    add_heading_styled(doc, "8. Módulos funcionales", 1)
    add_table(
        doc,
        ["Módulo", "Descripción", "Prioridad"],
        [
            ["Catálogo", "Listado, filtros, estados de stock", "Alta"],
            ["Ficha vehículo", "Especificaciones, galería, CTA contacto", "Alta"],
            ["Tour 360°", "Visualización interactiva optimizada para móvil", "Alta"],
            ["Admin", "Alta/edición de autos y medios", "Alta"],
            ["Contacto", "WhatsApp / formulario de interés", "Alta"],
            ["WebPay", "Pago/reserva online (según contrato)", "Media"],
            ["Financiamiento", "Derivación Forum u otra", "Media"],
            ["SEO / analytics", "Indexación y medición básica", "Media"],
        ],
    )

    add_heading_styled(doc, "9. Costos del proyecto (referenciales)", 1)
    add_heading_styled(doc, "9.1. Costos de infraestructura (recurrentes, Empresa)", 2)
    add_table(
        doc,
        ["Ítem", "Proveedor", "Valor referencial", "Periodicidad"],
        [
            ["Dominio .cl", "NIC Chile", "$9.990 CLP", "Anual"],
            ["Hosting", "Vercel Pro", "~US$20", "Mensual"],
            ["Base de datos + storage", "Supabase Pro", "~US$25", "Mensual"],
            ["Total cloud típico", "Vercel + Supabase", "~US$45 (~$40–50 mil CLP)", "Mensual"],
        ],
    )
    add_para(
        doc,
        "Tipos de cambio y tarifas de proveedores pueden variar. Valores a confirmar al contratar.",
        size=9,
        color=MUTED,
        italic=True,
    )

    add_heading_styled(doc, "9.2. Costos / gestiones comerciales (Empresa)", 2)
    add_table(
        doc,
        ["Ítem", "Observación"],
        [
            ["WebPay / Transbank", "Comisiones y cargos según contrato Transbank (no incluidos aquí)"],
            ["Forum / financieras", "Según acuerdo comercial dealer"],
            ["Producción audiovisual", "Si se externaliza fotógrafo/videógrafo; o costo interno de tiempo"],
            ["Mantención post-entrega", "Acuerdo aparte si se requiere soporte continuo"],
        ],
    )

    add_heading_styled(doc, "9.3. Esfuerzo de desarrollo", 2)
    add_para(
        doc,
        "El desarrollo se estima en un horizonte de 12 semanas de trabajo efectivo, sujeto a la entrega "
        "oportuna de insumos (fotos, decisiones de dominio/hosting/DB y credenciales de terceros).",
        align="justify",
    )

    add_heading_styled(doc, "10. Cronograma de producción (3 meses)", 1)
    add_table(
        doc,
        ["Fase", "Semanas", "Entregables"],
        [
            ["1. Descubrimiento y setup", "1–2", "Alcance firmado; cuentas Empresa; dominio; Supabase/Vercel; planilla stock"],
            ["2. Fundación técnica", "3–4", "Modelo de datos; auth admin; migración desde prototipo; CI/CD"],
            ["3. Catálogo y admin", "5–7", "CRUD vehículos; carga de medios; publicación en staging"],
            ["4. Tour 360° y performance", "8–9", "Pipeline video→frames; optimización móvil; pruebas en 4G"],
            ["5. Integraciones comerciales", "10–11", "Contacto/WhatsApp; preparación WebPay; canal Forum (según avance Empresa)"],
            ["6. Go-live y cierre", "12", "Dominio en producción; capacitación; acta de entrega; desligue de accesos personales"],
        ],
    )
    add_para(
        doc,
        "Hito crítico: sin fotos/videos ni decisión de dominio/hosting/DB en las primeras 2 semanas, "
        "el cronograma se desplaza automáticamente.",
        bold=True,
        size=10,
        color=NAVY,
    )

    add_heading_styled(doc, "11. Riesgos y mitigaciones", 1)
    add_table(
        doc,
        ["Riesgo", "Impacto", "Mitigación"],
        [
            ["Demora en entrega de medios", "Alto", "Documento de solicitud + checklist + Drive"],
            ["rgmotors.cl no disponible", "Medio", "Empresa propone 3 nombres .cl"],
            ["Cuentas cloud a nombre personal", "Alto", "Alta obligatoria con correo corporativo"],
            ["360° lento en móvil", "Alto", "Optimizar frames, lazy-load, CDN"],
            ["WebPay sin contrato a tiempo", "Medio", "v1 con contacto; WebPay como fase condicionada"],
            ["Forum sin API/clearance", "Medio", "Link/WhatsApp mientras se gestiona canal oficial"],
            ["Salida del desarrollador", "Alto", "Admin + docs + cuentas Empresa + capacitación"],
        ],
    )

    add_heading_styled(doc, "12. Plan de desligue / continuidad operativa", 1)
    for t in [
        "Todas las cuentas (NIC, GitHub org, Vercel, Supabase, Transbank) a nombre de la Empresa.",
        "Entrega de manual de operación del admin y credenciales al responsable interno.",
        "Repositorio en organización GitHub de la Empresa; retiro de accesos personales del desarrollador.",
        "Acta de recepción conforme firmada.",
        "Soporte posterior solo si se firma mantención (opcional).",
    ]:
        add_para(doc, f"• {t}", space_after=4)

    add_heading_styled(doc, "13. Criterios de aceptación (v1)", 1)
    for t in [
        "Sitio público accesible por dominio .cl con HTTPS.",
        "Al menos un set real de vehículos de Puerto Montt publicados con fotos.",
        "Al menos un tour 360° funcional y usable en móvil.",
        "Admin permite crear/editar/ocultar un vehículo sin intervención del desarrollador.",
        "Documento de capacitación entregado y sesión realizada.",
        "Cuentas cloud y dominio bajo control de la Empresa.",
    ]:
        add_para(doc, f"• {t}", space_after=4)

    add_heading_styled(doc, "14. Aspectos legales y de datos", 1)
    for t in [
        "Publicar políticas de privacidad y términos en el sitio (requisito también para WebPay).",
        "No publicar patentes ni datos sensibles de clientes sin base legal/autorización.",
        "Las imágenes/videos entregados se asumen con derechos de uso comercial por parte de la Empresa.",
        "Cumplir buenas prácticas de seguridad (HTTPS, control de acceso al admin, secretos fuera del código).",
    ]:
        add_para(doc, f"• {t}", space_after=4)

    add_heading_styled(doc, "15. Anexos", 1)
    for t in [
        "Anexo A: Documento RG-PM-SOL-001 (Solicitud de decisiones y accesos).",
        "Anexo B: Checklist de captura foto/video por vehículo.",
        "Anexo C: Glosario (dominio, hosting, base de datos, deploy, WebPay, staging/producción).",
    ]:
        add_para(doc, f"• {t}", space_after=4)

    add_heading_styled(doc, "16. Aprobación del documento de proyecto", 1)
    add_table(
        doc,
        ["Rol", "Nombre", "Firma", "Fecha"],
        [
            ["Elaborado por (responsable técnico)", "", "", ""],
            ["Revisado por (contraparte Empresa)", "", "", ""],
            ["Aprobado por (gerencia / administración)", "", "", ""],
        ],
        col_widths=[5.5, 4, 3.5, 3],
    )
    add_para(
        doc,
        "Este documento constituye la línea base del proyecto. Cualquier cambio de alcance deberá registrarse por escrito.",
        size=9,
        align="justify",
        color=MUTED,
        italic=True,
    )

    path = os.path.join(OUT, "02-Documento-Proyecto-Completo-RGMotors-PuertoMontt.docx")
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# DOC 3
# ---------------------------------------------------------------------------
def build_checklist():
    doc = Document()
    setup_doc(doc)
    add_footer(doc, "RG-PM-ANX-B")

    cover_page(
        doc,
        doc_type="Anexo operativo",
        title="Checklist de captura\npor vehículo",
        subtitle="Guía práctica para fotografías, video 360°\ny datos comerciales del stock",
        code="RG-PM-ANX-B",
    )

    toc_page(
        doc,
        [
            ("1.", "Identificación del vehículo", 1),
            ("2.", "Fotografías requeridas", 1),
            ("3.", "Video 360°", 1),
            ("4.", "Datos comerciales", 1),
            ("5.", "Conformidad de captura", 1),
        ],
        title="Índice",
    )

    add_heading_styled(doc, "1. Identificación del vehículo", 1)
    add_para(
        doc,
        "Vehículo: _______________________________  Patente interna: ____________  Fecha: ____/____/________",
        space_after=12,
    )

    add_heading_styled(doc, "2. Fotografías requeridas", 1)
    section_intro(doc, "Marque cada toma completada. Preferir JPG/PNG de al menos 1920 px.")
    for t in [
        "Frente",
        "Atras",
        "Lateral conductor",
        "Lateral copiloto",
        "3/4 delantero",
        "3/4 trasero",
        "Interior tablero",
        "Asientos traseros",
        "Maletero",
        "Motor",
        "Detalle llanta/neumático",
        "Foto hero (listado)",
    ]:
        add_para(doc, cb(t), space_after=4)

    add_heading_styled(doc, "3. Video 360°", 1)
    for t in [
        "Video de vuelta completa (30–60 s)",
        "Horizontal / trípode o paso firme",
        "Buena luz (evitar contraluz fuerte)",
        "Subido a Drive/OneDrive en calidad original (no WhatsApp)",
    ]:
        add_para(doc, cb(t), space_after=4)

    add_heading_styled(doc, "4. Datos comerciales", 1)
    for t in [
        "Precio actualizado",
        "Kilometraje",
        "Año / combustible / transmisión / color",
        "Estado: disponible / reservado / vendido",
    ]:
        add_para(doc, cb(t), space_after=4)

    add_heading_styled(doc, "5. Conformidad de captura", 1)
    add_para(doc, "Responsable captura: ________________________", space_after=10)
    add_para(doc, "Firma: ________________     Fecha: ____/____/________", space_after=10)
    add_para(
        doc,
        "Observaciones: ________________________________________________________________",
        space_after=8,
    )
    add_para(doc, "_" * 78, color=LINE)

    path = os.path.join(OUT, "03-Anexo-Checklist-Captura-Vehiculo.docx")
    doc.save(path)
    return path


if __name__ == "__main__":
    p1 = build_solicitud()
    p2 = build_proyecto()
    p3 = build_checklist()
    print("Generados:")
    print(p1)
    print(p2)
    print(p3)
